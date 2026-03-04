from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from typing import List


def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


class WordExporter:
    @staticmethod
    def _new_document() -> Document:
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        return doc

    def export(self, tables: List[List[List[str]]], output_path: str) -> None:
        doc = self._new_document()

        for i, table in enumerate(tables):
            if i > 0:
                doc.add_paragraph()

            if not table:
                continue

            cols = max(len(row) for row in table)
            t = doc.add_table(rows=len(table), cols=cols)
            t.style = "Table Grid"

            for r_idx, row in enumerate(table):
                cells = t.rows[r_idx].cells
                for c_idx in range(cols):
                    val = row[c_idx] if c_idx < len(row) else ""
                    cells[c_idx].text = val
                    if r_idx == 0:
                        _set_cell_bg(cells[c_idx], "2E86AB")
                        run = cells[c_idx].paragraphs[0].runs
                        if run:
                            run[0].bold = True
                            run[0].font.color.rgb = None
                        para = cells[c_idx].paragraphs[0]
                        for r in para.runs:
                            r.font.size = Pt(10)
                            r.bold = True

        doc.save(output_path)

    def export_text(self, text_pages: List[str], output_path: str) -> None:
        doc = self._new_document()

        wrote_text = False
        for i, page_text in enumerate(text_pages, 1):
            if i > 1:
                doc.add_page_break()

            doc.add_heading(f"Page {i}", level=2)
            lines = [line.strip() for line in (page_text or "").splitlines() if line.strip()]
            if not lines:
                doc.add_paragraph("")
                continue

            wrote_text = True
            for line in lines:
                doc.add_paragraph(line)

        if not text_pages or not wrote_text:
            doc.add_paragraph("No extractable text found in this PDF.")

        doc.save(output_path)
